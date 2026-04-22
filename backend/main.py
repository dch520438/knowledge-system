# -*- coding: utf-8 -*-
"""
FastAPI主应用模块
包含知识库、写作文档、问答记录的CRUD路由和全局搜索功能
以及批量导入导出、智能问答等扩展功能
"""

from typing import Optional, List
from fastapi import FastAPI, Depends, HTTPException, Query, UploadFile, File, Request
from fastapi.responses import StreamingResponse, Response, FileResponse
from sqlalchemy.orm import Session
from sqlalchemy import or_
from urllib.parse import quote as url_quote, urljoin, urlparse, urlencode
import io
import json
import csv
import re
import os

import httpx
from bs4 import BeautifulSoup

from docx import Document
from docx.shared import Cm
import pdfplumber

from database import engine, Base, get_db, SessionLocal
from models import KnowledgeItem, WritingDocument, QARecord, ProofreadRule, LLMConfig
from schemas import (
    KnowledgeItemCreate, KnowledgeItemResponse, KnowledgeItemUpdate,
    WritingDocumentCreate, WritingDocumentResponse, WritingDocumentUpdate,
    QARecordCreate, QARecordResponse, QARecordUpdate,
    SearchResult,
    BatchKnowledgeCreate,
    SmartAnswerRequest,
    ProofreadRuleCreate, ProofreadRuleResponse,
    LLMConfigCreate, LLMConfigUpdate, LLMConfigResponse,
    LLMChatRequest, LLMChatResponse,
    LLMQARequest, LLMWritingRequest, LLMKnowledgeRequest, LLMProofreadRequest,
)

# 创建FastAPI应用实例
app = FastAPI(
    title="知识管理系统 API",
    description="提供知识库管理、写作文档管理和问答记录管理的RESTful API",
    version="1.0.0",
)

# 添加CORS中间件，允许所有来源访问（开发模式）
from fastapi.middleware.cors import CORSMiddleware

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ==================== 启动事件：自动创建数据库表 ====================

@app.on_event("startup")
def startup_event():
    """应用启动时自动创建所有数据库表"""
    Base.metadata.create_all(bind=engine)
    # 初始化预置核稿规则
    init_proofread_rules()


# ==================== 知识库路由 ====================

@app.post("/api/knowledge", response_model=KnowledgeItemResponse, summary="创建知识条目")
def create_knowledge_item(item: KnowledgeItemCreate, db: Session = Depends(get_db)):
    """
    创建一个新的知识条目
    - **title**: 知识标题（必填）
    - **content**: 知识内容（必填）
    - **category**: 分类（可选）
    - **tags**: 标签，逗号分隔（可选）
    """
    db_item = KnowledgeItem(**item.model_dump())
    db.add(db_item)
    db.commit()
    db.refresh(db_item)
    return db_item


@app.post("/api/knowledge/search", response_model=List[KnowledgeItemResponse], summary="搜索知识条目")
def search_knowledge_items(
    search_data: dict,
    db: Session = Depends(get_db),
):
    """
    通过POST请求搜索知识条目，避免URL编码问题
    - **keyword**: 搜索关键词
    - **category**: 按分类筛选
    """
    keyword = search_data.get("keyword", "")
    category = search_data.get("category", "")

    query = db.query(KnowledgeItem)
    if category:
        query = query.filter(KnowledgeItem.category == category)
    if keyword:
        query = query.filter(
            or_(
                KnowledgeItem.title.contains(keyword),
                KnowledgeItem.content.contains(keyword),
                KnowledgeItem.category.contains(keyword),
                KnowledgeItem.tags.contains(keyword),
            )
        )
    items = query.order_by(KnowledgeItem.created_at.desc()).all()
    return items


@app.get("/api/knowledge", response_model=List[KnowledgeItemResponse], summary="获取知识列表")
def get_knowledge_items(
    category: Optional[str] = Query(None, description="按分类筛选"),
    db: Session = Depends(get_db),
):
    """
    获取知识条目列表，支持按分类筛选
    - **category**: 按分类筛选
    """
    query = db.query(KnowledgeItem)
    if category:
        query = query.filter(KnowledgeItem.category == category)
    items = query.order_by(KnowledgeItem.created_at.desc()).all()
    return items


@app.get("/api/knowledge/categories", response_model=List[str], summary="获取所有分类列表")
def get_categories(db: Session = Depends(get_db)):
    """获取所有不重复的分类列表"""
    categories = db.query(KnowledgeItem.category).distinct().filter(KnowledgeItem.category.isnot(None)).all()
    return [c[0] for c in categories]


@app.get("/api/knowledge/export", summary="导出知识库")
def export_knowledge(
    format: str = Query("json", description="导出格式：json, docx, csv"),
    db: Session = Depends(get_db),
):
    """
    导出知识库中所有知识条目
    - **format**: 导出格式，支持 json / docx / csv
    - **JSON**: 返回所有知识的JSON数组
    - **DOCX**: 生成Word文档，每个知识一个标题+内容段落
    - **CSV**: 生成CSV文件
    """
    items = db.query(KnowledgeItem).order_by(KnowledgeItem.created_at.desc()).all()

    if format == "json":
        # JSON格式导出
        data = []
        for item in items:
            data.append({
                "title": item.title,
                "content": item.content,
                "category": item.category,
                "tags": item.tags,
            })
        json_bytes = json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")
        return StreamingResponse(
            io.BytesIO(json_bytes),
            media_type="application/json",
            headers={"Content-Disposition": "attachment; filename*=UTF-8''knowledge_export.json"},
        )

    elif format == "docx":
        # DOCX格式导出（标记模板格式，可直接重新导入）
        doc = Document()
        for i, item in enumerate(items):
            doc.add_paragraph(f"标题：{item.title}")
            if item.category:
                doc.add_paragraph(f"分类：{item.category}")
            if item.tags:
                doc.add_paragraph(f"标签：{item.tags}")
            doc.add_paragraph(f"内容：{item.content}")
            if i < len(items) - 1:
                doc.add_paragraph("")
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename*=UTF-8''knowledge_export.docx"},
        )

    elif format == "csv":
        # CSV格式导出
        buffer = io.StringIO()
        writer = csv.writer(buffer)
        # 写入标题行
        writer.writerow(["标题", "分类", "标签", "内容"])
        # 写入数据行
        for item in items:
            writer.writerow([
                item.title,
                item.category or "",
                item.tags or "",
                item.content,
            ])
        csv_bytes = buffer.getvalue().encode("utf-8-sig")  # 使用utf-8-sig以支持Excel打开
        return StreamingResponse(
            io.BytesIO(csv_bytes),
            media_type="text/csv",
            headers={"Content-Disposition": "attachment; filename*=UTF-8''knowledge_export.csv"},
        )

    else:
        raise HTTPException(status_code=400, detail=f"不支持的导出格式: {format}，请使用 json / docx / csv")


# ==================== 批量操作路由（必须在 {item_id} 路由之前） ====================

@app.post("/api/knowledge/batch-delete", summary="批量删除知识")
def batch_delete_knowledge(
    data: dict,
    db: Session = Depends(get_db),
):
    """批量删除知识条目
    - **ids**: 要删除的知识条目ID列表
    """
    ids = data.get("ids", [])
    if not ids:
        raise HTTPException(status_code=400, detail="请提供要删除的ID列表")
    count = db.query(KnowledgeItem).filter(KnowledgeItem.id.in_(ids)).delete(synchronize_session=False)
    db.commit()
    return {"message": f"已删除 {count} 条知识", "count": count}


@app.put("/api/knowledge/batch-category", summary="批量更改分类")
def batch_update_category(
    data: dict,
    db: Session = Depends(get_db),
):
    """批量更改知识条目的分类
    - **ids**: 知识条目ID列表
    - **category**: 新的分类名称
    """
    ids = data.get("ids", [])
    category = data.get("category", "")
    if not ids:
        raise HTTPException(status_code=400, detail="请提供ID列表")
    if not category:
        raise HTTPException(status_code=400, detail="请提供新分类")
    count = db.query(KnowledgeItem).filter(KnowledgeItem.id.in_(ids)).update(
        {"category": category}, synchronize_session=False
    )
    db.commit()
    return {"message": f"已更新 {count} 条知识的分类", "count": count}


@app.put("/api/knowledge/batch-tags", summary="批量更改标签")
def batch_update_tags(
    data: dict,
    db: Session = Depends(get_db),
):
    """批量更改知识条目的标签
    - **ids**: 知识条目ID列表
    - **tags**: 新的标签（逗号分隔）
    - **mode**: 更新模式，"replace"替换/"append"追加/"remove"移除，默认"replace"
    """
    ids = data.get("ids", [])
    tags = data.get("tags", "")
    mode = data.get("mode", "replace")
    if not ids:
        raise HTTPException(status_code=400, detail="请提供ID列表")
    
    items = db.query(KnowledgeItem).filter(KnowledgeItem.id.in_(ids)).all()
    count = 0
    for item in items:
        existing_tags = [t.strip() for t in item.tags.split(",") if t.strip()] if item.tags else []
        new_tags = [t.strip() for t in tags.split(",") if t.strip()]
        
        if mode == "replace":
            final_tags = new_tags
        elif mode == "append":
            final_tags = list(set(existing_tags + new_tags))
        elif mode == "remove":
            final_tags = [t for t in existing_tags if t not in new_tags]
        else:
            final_tags = new_tags
        
        item.tags = ",".join(final_tags)
        count += 1
    
    db.commit()
    return {"message": f"已更新 {count} 条知识的标签", "count": count}


@app.get("/api/knowledge/{item_id}", response_model=KnowledgeItemResponse, summary="获取单个知识条目")
def get_knowledge_item(item_id: int, db: Session = Depends(get_db)):
    """根据ID获取单个知识条目"""
    db_item = db.query(KnowledgeItem).filter(KnowledgeItem.id == item_id).first()
    if db_item is None:
        raise HTTPException(status_code=404, detail="知识条目不存在")
    return db_item


@app.put("/api/knowledge/{item_id}", response_model=KnowledgeItemResponse, summary="更新知识条目")
def update_knowledge_item(item_id: int, item: KnowledgeItemUpdate, db: Session = Depends(get_db)):
    """
    更新指定ID的知识条目
    只更新请求中提供的字段，未提供的字段保持不变
    """
    db_item = db.query(KnowledgeItem).filter(KnowledgeItem.id == item_id).first()
    if db_item is None:
        raise HTTPException(status_code=404, detail="知识条目不存在")

    update_data = item.model_dump(exclude_unset=True)
    for key, value in update_data.items():
        setattr(db_item, key, value)

    db.commit()
    db.refresh(db_item)
    return db_item


@app.delete("/api/knowledge/{item_id}", summary="删除知识条目")
def delete_knowledge_item(item_id: int, db: Session = Depends(get_db)):
    """删除指定ID的知识条目"""
    db_item = db.query(KnowledgeItem).filter(KnowledgeItem.id == item_id).first()
    if db_item is None:
        raise HTTPException(status_code=404, detail="知识条目不存在")
    db.delete(db_item)
    db.commit()
    return {"message": "知识条目已删除", "id": item_id}


@app.post("/api/knowledge/deduplicate", summary="知识库排重")
def deduplicate_knowledge(db: Session = Depends(get_db)):
    """
    删除内容完全相同的知识条目，每个重复内容只保留最早创建的一条
    返回删除的数量
    """
    # 获取所有知识条目，按创建时间升序
    items = db.query(KnowledgeItem).order_by(KnowledgeItem.created_at.asc()).all()
    
    # 用字典记录已见过的内容
    seen_content = {}
    duplicates_to_delete = []
    
    for item in items:
        # 标准化内容用于比较（去除首尾空白）
        normalized_content = item.content.strip() if item.content else ""
        
        if normalized_content in seen_content:
            # 这是重复内容，标记删除
            duplicates_to_delete.append(item.id)
        else:
            # 首次见到此内容，记录
            seen_content[normalized_content] = item.id
    
    # 删除重复条目
    for item_id in duplicates_to_delete:
        db.query(KnowledgeItem).filter(KnowledgeItem.id == item_id).delete()
    
    db.commit()
    
    return {"removed_count": len(duplicates_to_delete), "message": f"成功删除 {len(duplicates_to_delete)} 条重复知识"}


# ==================== 知识库批量导入/导出路由 ====================

@app.post("/api/knowledge/batch", response_model=List[KnowledgeItemResponse], summary="批量创建知识条目")
def batch_create_knowledge_items(batch: BatchKnowledgeCreate, db: Session = Depends(get_db)):
    """
    批量创建知识条目
    - **items**: 知识条目数组，每个条目包含 title, content, category, tags
    """
    created_items = []
    for item_data in batch.items:
        db_item = KnowledgeItem(**item_data.model_dump())
        db.add(db_item)
        created_items.append(db_item)
    db.commit()
    for item in created_items:
        db.refresh(item)
    return created_items


@app.post("/api/knowledge/import", response_model=List[KnowledgeItemResponse], summary="从文件导入知识")
async def import_knowledge_from_file(
    file: UploadFile = File(..., description="上传文件，支持 .json / .docx / .pdf / .txt / .csv"),
    db: Session = Depends(get_db),
):
    """
    从文件导入知识条目，支持多种文件格式：
    - **JSON**: 解析数组 [{title, content, category, tags}]
    - **DOCX**: 每个段落作为一条知识（标题取前20字）
    - **PDF**: 每页作为一条知识（标题取前20字）
    - **TXT**: 按空行分割，每段作为一条知识
    - **CSV**: 第一行标题列和内容列，后续行数据
    """
    # 读取文件内容
    content_bytes = await file.read()

    # 获取文件扩展名
    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    knowledge_list = []  # 存储解析后的知识条目 {title, content, category, tags}

    if ext == "json":
        # JSON格式：解析数组 [{title, content, category, tags}]
        try:
            text = content_bytes.decode("utf-8")
            data = json.loads(text)
            if not isinstance(data, list):
                raise HTTPException(status_code=400, detail="JSON文件内容必须是数组格式")
            for item in data:
                title = item.get("title", "").strip()
                content = item.get("content", "").strip()
                if not title or not content:
                    continue
                knowledge_list.append({
                    "title": title[:255],
                    "content": content,
                    "category": item.get("category"),
                    "tags": item.get("tags", ""),
                })
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="JSON文件解析失败")

    elif ext == "docx":
        # DOCX格式：支持标题样式和标记格式
        try:
            doc = Document(io.BytesIO(content_bytes))
            
            # 检测是否包含标记格式（标题：分类：等）
            full_text = "\n".join([p.text for p in doc.paragraphs])
            if re.search(r'标题[：:]', full_text):
                # 标记模板格式
                entries = re.split(r'\n\s*(?=标题[：:])', full_text.strip())
                for entry in entries:
                    entry = entry.strip()
                    if not entry:
                        continue
                    title_match = re.search(r'标题[：:]\s*(.+)', entry)
                    category_match = re.search(r'分类[：:]\s*(.+)', entry)
                    tags_match = re.search(r'标签[：:]\s*(.+)', entry)
                    content_match = re.search(r'内容[：:]\s*(.+)', entry, re.DOTALL)
                    
                    title = title_match.group(1).strip() if title_match else ""
                    content = content_match.group(1).strip() if content_match else ""
                    if not title or not content:
                        continue
                    knowledge_list.append({
                        "title": title[:255],
                        "content": content,
                        "category": category_match.group(1).strip() if category_match else None,
                        "tags": tags_match.group(1).strip() if tags_match else "",
                    })
            else:
                # 普通格式：每个段落作为一条知识
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    title = text[:20] + ("..." if len(text) > 20 else "")
                    knowledge_list.append({
                        "title": title,
                        "content": text,
                        "category": None,
                        "tags": "",
                    })
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"DOCX文件解析失败: {str(e)}")

    elif ext == "pdf":
        # PDF格式：每页作为一条知识（标题取前20字）
        try:
            with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
                for page in pdf.pages:
                    text = (page.extract_text() or "").strip()
                    if not text:
                        continue
                    title = text[:20] + ("..." if len(text) > 20 else "")
                    knowledge_list.append({
                        "title": title,
                        "content": text,
                        "category": None,
                        "tags": "",
                    })
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"PDF文件解析失败: {str(e)}")

    elif ext == "txt":
        # TXT格式：支持标记模板和空行分割两种格式
        try:
            text = content_bytes.decode("utf-8")
            
            # 检测是否为标记模板格式（包含"标题："或"标题:"）
            if re.search(r'标题[：:]', text):
                # 标记模板格式：按"标题："分割为多个知识条目
                entries = re.split(r'\n\s*(?=标题[：:])', text.strip())
                for entry in entries:
                    entry = entry.strip()
                    if not entry:
                        continue
                    title_match = re.search(r'标题[：:]\s*(.+)', entry)
                    category_match = re.search(r'分类[：:]\s*(.+)', entry)
                    tags_match = re.search(r'标签[：:]\s*(.+)', entry)
                    content_match = re.search(r'内容[：:]\s*(.+)', entry, re.DOTALL)
                    
                    title = title_match.group(1).strip() if title_match else ""
                    content = content_match.group(1).strip() if content_match else ""
                    if not title or not content:
                        continue
                    knowledge_list.append({
                        "title": title[:255],
                        "content": content,
                        "category": category_match.group(1).strip() if category_match else None,
                        "tags": tags_match.group(1).strip() if tags_match else "",
                    })
            else:
                # 普通格式：按空行分割段落
                paragraphs = re.split(r'\n\s*\n', text)
                for para in paragraphs:
                    para_text = para.strip()
                    if not para_text:
                        continue
                    para_text = re.sub(r'\n', ' ', para_text).strip()
                    title = para_text[:20] + ("..." if len(para_text) > 20 else "")
                    knowledge_list.append({
                        "title": title,
                        "content": para_text,
                        "category": None,
                        "tags": "",
                    })
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"TXT文件解析失败: {str(e)}")

    elif ext == "csv":
        # CSV格式：第一行标题列和内容列，后续行数据
        try:
            text = content_bytes.decode("utf-8")
            reader = csv.reader(io.StringIO(text))
            rows = list(reader)
            if len(rows) < 2:
                raise HTTPException(status_code=400, detail="CSV文件至少需要标题行和一行数据")
            # 读取标题行，查找标题列和内容列
            header = [h.strip().lower() for h in rows[0]]
            title_idx = None
            content_idx = None
            category_idx = None
            tags_idx = None
            for i, col_name in enumerate(header):
                if col_name in ("title", "标题", "名称"):
                    title_idx = i
                elif col_name in ("content", "内容", "正文"):
                    content_idx = i
                elif col_name in ("category", "分类", "类别"):
                    category_idx = i
                elif col_name in ("tags", "标签"):
                    tags_idx = i
            # 如果没有找到标题列和内容列，默认第一列为标题，第二列为内容
            if title_idx is None:
                title_idx = 0
            if content_idx is None:
                content_idx = 1 if len(header) > 1 else 0
            for row in rows[1:]:
                if not row:
                    continue
                title = row[title_idx].strip() if title_idx < len(row) else ""
                content = row[content_idx].strip() if content_idx < len(row) else ""
                if not title or not content:
                    continue
                knowledge_list.append({
                    "title": title[:255],
                    "content": content,
                    "category": row[category_idx].strip() if category_idx is not None and category_idx < len(row) else None,
                    "tags": row[tags_idx].strip() if tags_idx is not None and tags_idx < len(row) else "",
                })
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"CSV文件解析失败: {str(e)}")

    else:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: .{ext}，请上传 .json / .docx / .pdf / .txt / .csv 格式的文件"
        )

    if not knowledge_list:
        raise HTTPException(status_code=400, detail="未能从文件中解析出有效的知识条目")

    # 批量写入数据库
    created_items = []
    for item_data in knowledge_list:
        db_item = KnowledgeItem(**item_data)
        db.add(db_item)
        created_items.append(db_item)
    db.commit()
    for item in created_items:
        db.refresh(item)

    return created_items


@app.post("/api/knowledge/from-writing", response_model=KnowledgeItemResponse, summary="从写作文稿创建知识")
def create_knowledge_from_writing(
    data: dict,
    db: Session = Depends(get_db),
):
    """
    从指定写作文档中提取内容创建知识条目
    - **doc_id**: 写作文档ID
    - **title**: 知识标题
    - **content**: 知识内容（如不提供则使用写作文档的内容）
    - **category**: 分类（可选）
    - **tags**: 标签（可选）
    """
    doc_id = data.get("doc_id")
    if not doc_id:
        raise HTTPException(status_code=400, detail="doc_id 为必填参数")

    # 查找写作文档
    writing_doc = db.query(WritingDocument).filter(WritingDocument.id == doc_id).first()
    if writing_doc is None:
        raise HTTPException(status_code=404, detail="写作文档不存在")

    # 获取参数
    title = data.get("title", "").strip()
    content = data.get("content", "").strip()
    category = data.get("category")
    tags = data.get("tags", "")

    # 如果未提供标题，使用写作文档标题
    if not title:
        title = writing_doc.title
    # 如果未提供内容，使用写作文档内容
    if not content:
        content = writing_doc.content

    if not title or not content:
        raise HTTPException(status_code=400, detail="标题和内容不能同时为空")

    # 创建知识条目
    db_item = KnowledgeItem(
        title=title[:255],
        content=content,
        category=category,
        tags=tags,
    )
    db.add(db_item)
    db.commit()
    db.refresh(db_item)
    return db_item


# ==================== 写作文档路由 ====================

@app.post("/api/writing", response_model=WritingDocumentResponse, summary="创建写作文档")
def create_writing_document(doc: WritingDocumentCreate, db: Session = Depends(get_db)):
    """
    创建一个新的写作文档
    - **title**: 文档标题（必填）
    - **content**: 文档内容
    - **referenced_knowledge_ids**: 引用的知识条目ID，逗号分隔
    - **status**: 状态，draft-草稿，published-已发布
    """
    db_doc = WritingDocument(**doc.model_dump())
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)
    return db_doc


@app.get("/api/writing", response_model=List[WritingDocumentResponse], summary="获取文档列表")
def get_writing_documents(
    status: Optional[str] = Query(None, description="按状态筛选：draft/published"),
    db: Session = Depends(get_db),
):
    """
    获取写作文档列表，支持按状态筛选
    - **status**: 按状态筛选（draft-草稿，published-已发布）
    """
    query = db.query(WritingDocument)
    if status:
        query = query.filter(WritingDocument.status == status)
    docs = query.order_by(WritingDocument.created_at.desc()).all()
    return docs


@app.get("/api/writing/{doc_id}", response_model=WritingDocumentResponse, summary="获取单个文档")
def get_writing_document(doc_id: int, db: Session = Depends(get_db)):
    """根据ID获取单个写作文档"""
    db_doc = db.query(WritingDocument).filter(WritingDocument.id == doc_id).first()
    if db_doc is None:
        raise HTTPException(status_code=404, detail="写作文档不存在")
    return db_doc


@app.put("/api/writing/{doc_id}", response_model=WritingDocumentResponse, summary="更新文档")
def update_writing_document(doc_id: int, doc: WritingDocumentUpdate, db: Session = Depends(get_db)):
    """
    更新指定ID的写作文档
    只更新请求中提供的字段，未提供的字段保持不变
    """
    db_doc = db.query(WritingDocument).filter(WritingDocument.id == doc_id).first()
    if db_doc is None:
        raise HTTPException(status_code=404, detail="写作文档不存在")

    update_data = doc.model_dump(exclude_unset=True)
    for key, value in update_data.items():
        setattr(db_doc, key, value)

    db.commit()
    db.refresh(db_doc)
    return db_doc


@app.delete("/api/writing/{doc_id}", summary="删除文档")
def delete_writing_document(doc_id: int, db: Session = Depends(get_db)):
    """删除指定ID的写作文档"""
    db_doc = db.query(WritingDocument).filter(WritingDocument.id == doc_id).first()
    if db_doc is None:
        raise HTTPException(status_code=404, detail="写作文档不存在")
    db.delete(db_doc)
    db.commit()
    return {"message": "写作文档已删除", "id": doc_id}


# ==================== 写作文档导入/导出路由 ====================

def html_to_docx(html_content: str) -> Document:
    """将 Quill HTML 内容转换为 python-docx Document，保留格式"""
    doc = Document()
    if not html_content or not html_content.strip():
        return doc
    
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 遍历所有顶层元素
    for element in soup.children:
        if isinstance(element, str):
            text = element.strip()
            if text:
                doc.add_paragraph(text)
            continue
        
        tag = element.name
        if tag is None:
            continue
        
        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            level = int(tag[1])
            text = element.get_text()
            if text.strip():
                doc.add_heading(text.strip(), level=level)
        
        elif tag == 'p':
            para = doc.add_paragraph()
            # 处理段落中的内联格式
            _add_formatted_runs(para, element)
        
        elif tag in ('ul', 'ol'):
            list_type = 'ol' if tag == 'ol' else 'ul'
            for li in element.find_all('li', recursive=False):
                para = doc.add_paragraph(style='List Bullet' if list_type == 'ul' else 'List Number')
                _add_formatted_runs(para, li)
        
        elif tag == 'blockquote':
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1.5)
            _add_formatted_runs(para, element)
        
        elif tag == 'br':
            doc.add_paragraph('')
        
        elif tag in ('div', 'section'):
            # 容器元素，递归处理子元素
            for child in element.children:
                if isinstance(child, str):
                    text = child.strip()
                    if text:
                        doc.add_paragraph(text)
                elif child.name == 'p':
                    para = doc.add_paragraph()
                    _add_formatted_runs(para, child)
                elif child.name == 'br':
                    doc.add_paragraph('')
                elif child.name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                    level = int(child.name[1])
                    text = child.get_text()
                    if text.strip():
                        doc.add_heading(text.strip(), level=level)
                elif child.name in ('ul', 'ol'):
                    list_type = 'ol' if child.name == 'ol' else 'ul'
                    for li in child.find_all('li', recursive=False):
                        para = doc.add_paragraph(style='List Bullet' if list_type == 'ul' else 'List Number')
                        _add_formatted_runs(para, li)
                else:
                    text = child.get_text().strip()
                    if text:
                        doc.add_paragraph(text)
        
        else:
            text = element.get_text().strip()
            if text:
                doc.add_paragraph(text)
    
    return doc


def _add_formatted_runs(paragraph, element):
    """将 HTML 元素中的内联格式（加粗、斜体、下划线等）添加到 docx 段落"""
    from docx.shared import Pt
    
    def process_node(node):
        if isinstance(node, str):
            text = node
            if text:
                paragraph.add_run(text)
            return
        
        tag = node.name
        text = node.get_text()
        if not text:
            return
        
        run = paragraph.add_run(text)
        
        # 检查父元素链中的格式标记
        parent = node.parent
        while parent and parent.name:
            if parent.name in ('strong', 'b'):
                run.bold = True
            elif parent.name in ('em', 'i'):
                run.italic = True
            elif parent.name == 'u':
                run.underline = True
            elif parent.name == 's' or parent.name == 'del':
                run.strike = True
            elif parent.name == 'sub':
                run.font.subscript = True
            elif parent.name == 'sup':
                run.font.superscript = True
            parent = parent.parent
        
        # 也检查当前节点
        if tag in ('strong', 'b'):
            run.bold = True
        elif tag in ('em', 'i'):
            run.italic = True
        elif tag == 'u':
            run.underline = True
        elif tag == 's' or tag == 'del':
            run.strike = True
    
    for child in element.children:
        process_node(child)


def docx_to_html(doc: Document) -> str:
    """将 python-docx Document 转换为 Quill 兼容的 HTML，保留格式"""
    html_parts = []
    
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        text = para.text
        
        if not text.strip():
            continue
        
        # 检测标题级别
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.split()[-1])
                html_parts.append(f'<h{level}>{_runs_to_html(para.runs)}</h{level}>')
                continue
            except (ValueError, IndexError):
                pass
        
        # 检测列表
        if 'List' in style_name:
            if 'Number' in style_name or '2' in style_name:
                html_parts.append(f'<li>{_runs_to_html(para.runs)}</li>')
            else:
                html_parts.append(f'<li>{_runs_to_html(para.runs)}</li>')
            continue
        
        # 普通段落
        html_parts.append(f'<p>{_runs_to_html(para.runs)}</p>')
    
    return '\n'.join(html_parts)


def _runs_to_html(runs) -> str:
    """将 docx runs 转换为带格式的 HTML"""
    if not runs:
        return ''
    
    parts = []
    for run in runs:
        text = run.text or ''
        if not text:
            continue
        
        # 转义 HTML 特殊字符
        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        if run.bold:
            text = f'<strong>{text}</strong>'
        if run.italic:
            text = f'<em>{text}</em>'
        if run.underline:
            text = f'<u>{text}</u>'
        try:
            if run.font.strike:
                text = f'<s>{text}</s>'
        except AttributeError:
            pass
        
        parts.append(text)
    
    return ''.join(parts)


@app.get("/api/writing/{doc_id}/export", summary="导出单个写作文档")
def export_writing_document(
    doc_id: int,
    format: str = Query("docx", description="导出格式：docx"),
    db: Session = Depends(get_db),
):
    """
    导出单个写作文档
    - **doc_id**: 文档ID
    - **format**: 导出格式，目前支持 docx
    """
    db_doc = db.query(WritingDocument).filter(WritingDocument.id == doc_id).first()
    if db_doc is None:
        raise HTTPException(status_code=404, detail="写作文档不存在")

    if format == "docx":
        # 生成Word文档，保留格式
        doc = Document()
        doc.add_heading(db_doc.title, level=1)
        if db_doc.content:
            content_doc = html_to_docx(db_doc.content)
            # 将内容文档的所有元素追加到主文档
            for element in content_doc.element.body:
                doc.element.body.append(element)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        # 使用文档标题作为文件名，替换非法字符
        safe_filename = re.sub(r'[\\/:*?"<>|]', '_', db_doc.title)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{url_quote(safe_filename)}.docx"},
        )
    else:
        raise HTTPException(status_code=400, detail=f"不支持的导出格式: {format}，目前仅支持 docx")


@app.post("/api/writing/import", response_model=WritingDocumentResponse, summary="从文件导入文档")
async def import_writing_document(
    file: UploadFile = File(..., description="上传文件，支持 .docx / .txt"),
    db: Session = Depends(get_db),
):
    """
    从文件导入写作文档，支持多种文件格式：
    - **DOCX**: 读取内容创建新文档
    - **TXT**: 读取内容创建新文档
    """
    content_bytes = await file.read()

    # 获取文件扩展名
    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    doc_title = ""
    doc_content = ""

    if ext == "docx":
        # DOCX格式：读取内容，保留格式
        try:
            doc = Document(io.BytesIO(content_bytes))
            if not doc.paragraphs:
                raise HTTPException(status_code=400, detail="DOCX文件中没有有效内容")
            # 使用第一段作为标题
            doc_title = doc.paragraphs[0].text.strip()[:255]
            if not doc_title:
                doc_title = filename.rsplit('.', 1)[0][:255]
            # 将所有段落转为HTML（保留格式）
            doc_content = docx_to_html(doc)
            # 去掉标题段落，只保留内容
            soup = BeautifulSoup(doc_content, 'html.parser')
            first_p = soup.find(['p', 'h1', 'h2', 'h3'])
            if first_p:
                first_p.decompose()
            doc_content = str(soup).strip()
            if not doc_content:
                doc_content = '<p></p>'
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"DOCX文件解析失败: {str(e)}")

    elif ext == "txt":
        # TXT格式：读取内容，转为HTML
        try:
            text = content_bytes.decode("utf-8")
            lines = [line.strip() for line in text.split("\n") if line.strip()]
            if not lines:
                raise HTTPException(status_code=400, detail="TXT文件中没有有效内容")
            # 使用第一行作为标题，其余作为内容
            doc_title = lines[0][:255]
            # 将每行转为HTML段落
            content_lines = [f'<p>{line}</p>' for line in lines[1:]] if len(lines) > 1 else ['<p></p>']
            doc_content = "\n".join(content_lines)
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"TXT文件解析失败: {str(e)}")

    else:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: .{ext}，请上传 .docx / .txt 格式的文件"
        )

    # 创建写作文档
    db_doc = WritingDocument(
        title=doc_title,
        content=doc_content,
        status="draft",
    )
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)
    return db_doc


# ==================== 问答记录路由 ====================

@app.post("/api/qa", response_model=QARecordResponse, summary="创建问答记录")
def create_qa_record(qa: QARecordCreate, db: Session = Depends(get_db)):
    """
    创建一个新的问答记录
    - **question**: 问题（必填）
    - **answer**: 回答
    - **referenced_knowledge_ids**: 引用的知识条目ID，逗号分隔
    """
    db_qa = QARecord(**qa.model_dump())
    db.add(db_qa)
    db.commit()
    db.refresh(db_qa)
    return db_qa


@app.get("/api/qa", response_model=List[QARecordResponse], summary="获取问答列表")
def get_qa_records(db: Session = Depends(get_db)):
    """获取所有问答记录列表"""
    records = db.query(QARecord).order_by(QARecord.created_at.desc()).all()
    return records


@app.get("/api/qa/{qa_id}", response_model=QARecordResponse, summary="获取单个问答记录")
def get_qa_record(qa_id: int, db: Session = Depends(get_db)):
    """根据ID获取单个问答记录"""
    db_qa = db.query(QARecord).filter(QARecord.id == qa_id).first()
    if db_qa is None:
        raise HTTPException(status_code=404, detail="问答记录不存在")
    return db_qa


@app.put("/api/qa/{qa_id}", response_model=QARecordResponse, summary="更新问答记录")
def update_qa_record(qa_id: int, qa: QARecordUpdate, db: Session = Depends(get_db)):
    """
    更新指定ID的问答记录
    只更新请求中提供的字段，未提供的字段保持不变
    """
    db_qa = db.query(QARecord).filter(QARecord.id == qa_id).first()
    if db_qa is None:
        raise HTTPException(status_code=404, detail="问答记录不存在")

    update_data = qa.model_dump(exclude_unset=True)
    for key, value in update_data.items():
        setattr(db_qa, key, value)

    db.commit()
    db.refresh(db_qa)
    return db_qa


@app.delete("/api/qa/{qa_id}", summary="删除问答记录")
def delete_qa_record(qa_id: int, db: Session = Depends(get_db)):
    """删除指定ID的问答记录"""
    db_qa = db.query(QARecord).filter(QARecord.id == qa_id).first()
    if db_qa is None:
        raise HTTPException(status_code=404, detail="问答记录不存在")
    db.delete(db_qa)
    db.commit()
    return {"message": "问答记录已删除", "id": qa_id}


# ==================== 全局搜索路由 ====================

@app.get("/api/search", response_model=List[SearchResult], summary="全局搜索")
def global_search(
    q: str = Query(..., min_length=1, description="搜索关键词"),
    db: Session = Depends(get_db),
):
    """
    全局搜索知识库的标题和内容
    返回匹配的知识条目列表
    - **q**: 搜索关键词（必填）
    """
    results: List[SearchResult] = []

    # 在知识库中搜索标题和内容
    knowledge_items = db.query(KnowledgeItem).filter(
        or_(
            KnowledgeItem.title.contains(q),
            KnowledgeItem.content.contains(q),
            KnowledgeItem.category.contains(q),
            KnowledgeItem.tags.contains(q),
        )
    ).order_by(KnowledgeItem.created_at.desc()).all()

    for item in knowledge_items:
        results.append(SearchResult(
            type="knowledge",
            id=item.id,
            title=item.title,
            content=item.content[:200] if len(item.content) > 200 else item.content,
            category=item.category,
            created_at=item.created_at,
        ))

    return results


# ==================== 智能问答路由（关键词匹配版） ====================

@app.post("/api/qa/smart-answer", summary="智能问答")
def smart_answer(
    request: SmartAnswerRequest,
    db: Session = Depends(get_db),
):
    """
    基于关键词匹配的智能问答
    - **question**: 问题（必填）

    实现逻辑：
    1. 对问题进行分词（按空格、标点分割为关键词列表）
    2. 在知识库中搜索包含任一关键词的知识条目
    3. 对匹配结果按匹配关键词数量排序
    4. 取top 5结果
    5. 组织回答返回
    """
    question = request.question.strip()

    # 第一步：对问题进行分词（按空格、标点分割为关键词列表）
    # 使用正则表达式分割：按空格、中文标点、英文标点分割
    keywords = re.split(r'[\s,，。.!！?？;；:：、\(\)（）\[\]【】\""\']+/', question)
    # 过滤空字符串和过短的词（单字词可能匹配过多）
    keywords = [kw.strip() for kw in keywords if len(kw.strip()) >= 2]
    # 如果过滤后没有关键词，尝试使用原始分割结果（保留单字）
    if not keywords:
        keywords = [kw.strip() for kw in re.split(r'[\s,，。.!！?？;；:：、\(\)（）\[\]【】\""\']+/', question) if kw.strip()]

    if not keywords:
        return {
            "answer": "抱歉，知识库中未找到与您的问题相关的内容。建议您在知识库中添加相关知识，或尝试使用不同的关键词提问。",
            "sources": [],
            "keywords": [],
        }

    # 第二步：在知识库中搜索包含任一关键词的知识条目
    all_items = db.query(KnowledgeItem).all()

    # 第三步：对每个知识条目计算匹配的关键词数量
    scored_items = []
    for item in all_items:
        match_count = 0
        search_text = (item.title + " " + item.content + " " + (item.category or "") + " " + (item.tags or "")).lower()
        for kw in keywords:
            if kw.lower() in search_text:
                match_count += 1
        if match_count > 0:
            scored_items.append((item, match_count))

    # 第四步：按匹配关键词数量排序，取top 5
    scored_items.sort(key=lambda x: x[1], reverse=True)
    top_items = scored_items[:5]

    # 第五步：组织回答
    if not top_items:
        return {
            "answer": "抱歉，知识库中未找到与您的问题相关的内容。建议您在知识库中添加相关知识，或尝试使用不同的关键词提问。",
            "sources": [],
            "keywords": keywords,
        }

    answer_parts = ["根据知识库检索，找到以下相关内容：\n"]
    sources = []
    for item, score in top_items:
        answer_parts.append(f"\n【来源：{item.title}】\n{item.content}\n")
        sources.append({"id": item.id, "title": item.title})

    return {
        "answer": "".join(answer_parts),
        "sources": sources,
        "keywords": keywords,
    }


# ==================== 核稿规则路由 ====================

@app.get("/api/proofread/rules", response_model=List[ProofreadRuleResponse], summary="获取所有核稿规则")
def get_proofread_rules(db: Session = Depends(get_db)):
    """获取所有核稿规则列表"""
    return db.query(ProofreadRule).order_by(ProofreadRule.is_builtin.desc(), ProofreadRule.id.asc()).all()

@app.post("/api/proofread/rules", response_model=ProofreadRuleResponse, summary="创建核稿规则")
def create_proofread_rule(rule: ProofreadRuleCreate, db: Session = Depends(get_db)):
    """创建自定义核稿规则"""
    db_rule = ProofreadRule(**rule.model_dump())
    db.add(db_rule)
    db.commit()
    db.refresh(db_rule)
    return db_rule

@app.put("/api/proofread/rules/{rule_id}", response_model=ProofreadRuleResponse, summary="更新核稿规则")
def update_proofread_rule(rule_id: int, rule: ProofreadRuleCreate, db: Session = Depends(get_db)):
    """更新核稿规则"""
    db_rule = db.query(ProofreadRule).filter(ProofreadRule.id == rule_id).first()
    if not db_rule:
        raise HTTPException(status_code=404, detail="核稿规则不存在")
    for key, value in rule.model_dump().items():
        setattr(db_rule, key, value)
    db.commit()
    db.refresh(db_rule)
    return db_rule

@app.delete("/api/proofread/rules/{rule_id}", summary="删除核稿规则")
def delete_proofread_rule(rule_id: int, db: Session = Depends(get_db)):
    """删除核稿规则（不能删除预置规则）"""
    db_rule = db.query(ProofreadRule).filter(ProofreadRule.id == rule_id).first()
    if not db_rule:
        raise HTTPException(status_code=404, detail="核稿规则不存在")
    if db_rule.is_builtin:
        raise HTTPException(status_code=400, detail="预置规则不能删除")
    db.delete(db_rule)
    db.commit()
    return {"message": "删除成功"}

@app.post("/api/proofread/check", summary="核稿检查")
def proofread_check(
    body: dict,
    db: Session = Depends(get_db),
):
    """
    对文稿内容进行核稿检查
    返回所有匹配的错误列表，包含位置信息
    """
    content = body.get("content", "")
    rules = db.query(ProofreadRule).filter(ProofreadRule.enabled == True).all()

    results = []
    for rule in rules:
        try:
            pattern = re.compile(rule.pattern)
        except re.error:
            continue

        for match in pattern.finditer(content):
            results.append({
                "rule_id": rule.id,
                "rule_name": rule.name,
                "description": rule.description,
                "severity": rule.severity,
                "matched_text": match.group(),
                "start": match.start(),
                "end": match.end(),
                "line": content[:match.start()].count('\n') + 1,
            })

    # 按位置排序
    results.sort(key=lambda x: x["start"])

    return {"errors": results, "total": len(results)}


# ==================== 预置核稿规则初始化 ====================

def init_proofread_rules():
    """初始化预置核稿规则"""
    db = SessionLocal()
    try:
        existing = db.query(ProofreadRule).filter(ProofreadRule.is_builtin == True).count()
        if existing > 0:
            return

        builtin_rules = [
            {"name": "重复句号", "description": "连续出现多个句号", "pattern": "。{2,}", "severity": "error"},
            {"name": "重复逗号", "description": "连续出现多个逗号", "pattern": "，{2,}", "severity": "error"},
            {"name": "重复顿号", "description": "连续出现多个顿号", "pattern": "、{2,}", "severity": "error"},
            {"name": "多余空格", "description": "中文之间有多余空格", "pattern": "[\\u4e00-\\u9fff] +[\\u4e00-\\u9fff]", "severity": "warning"},
            {"name": "英文前后缺空格", "description": "中文与英文之间缺少空格", "pattern": "[\\u4e00-\\u9fff][a-zA-Z]|[a-zA-Z][\\u4e00-\\u9fff]", "severity": "info"},
            {"name": "重复的", "description": "出现重复的\"的\"", "pattern": "的的", "severity": "error"},
            {"name": "重复的了", "description": "出现重复的了", "pattern": "了了", "severity": "error"},
            {"name": "首行空格", "description": "段落首行有多余空格（建议使用缩进）", "pattern": "^ +", "severity": "info"},
            {"name": "全角数字", "description": "文中使用了全角数字", "pattern": "[０-９]", "severity": "warning"},
            {"name": "多余空行", "description": "连续出现多个空行", "pattern": "\n{3,}", "severity": "warning"},
        ]

        for r in builtin_rules:
            db_rule = ProofreadRule(
                name=r["name"],
                description=r["description"],
                pattern=r["pattern"],
                severity=r["severity"],
                is_builtin=True,
                enabled=True,
            )
            db.add(db_rule)
        db.commit()
    finally:
        db.close()


# ==================== 根路由 ====================



# ==================== 网页代理路由 ====================

@app.get("/api/proxy/web", summary="网页代理")
async def proxy_web(url: str = Query(..., description="要代理的网页URL")):
    """
    通过后端代理转发网页内容，去除X-Frame-Options等安全限制头，
    重写页面中的链接使其也通过代理加载，确保iframe可以正常嵌入显示。
    静态资源（图片、CSS、JS）直接使用原始URL，不经过代理，提升加载速度。
    """
    # 验证URL格式
    try:
        parsed = urlparse(url)
        if parsed.scheme not in ('http', 'https'):
            raise HTTPException(status_code=400, detail="仅支持 http/https 协议")
        if not parsed.netloc:
            raise HTTPException(status_code=400, detail="无效的URL")
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=400, detail="无效的URL")

    try:
        async with httpx.AsyncClient(
            follow_redirects=True,
            timeout=20.0,
            headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
                'Accept-Encoding': 'identity',
            }
        ) as client:
            response = await client.get(url)

        content_type = response.headers.get('content-type', 'text/html')

        # 非HTML内容（图片、CSS、JS、字体等），直接转发，不走HTML解析
        if 'text/html' not in content_type:
            return Response(
                content=response.content,
                media_type=content_type,
                headers={
                    "X-Frame-Options": "ALLOWALL",
                    "Content-Security-Policy": "frame-ancestors *",
                    "Access-Control-Allow-Origin": "*",
                    "Cache-Control": "public, max-age=86400",
                }
            )

        # HTML内容：重写导航链接，保留静态资源原始URL
        html_content = response.text
        soup = BeautifulSoup(html_content, 'html.parser')

        base_url = url
        parsed_base = urlparse(base_url)
        base_origin = f"{parsed_base.scheme}://{parsed_base.netloc}"

        def should_proxy(href):
            """判断链接是否需要代理（仅代理同站导航链接）"""
            if not href:
                return False
            href = href.strip()
            if href.startswith('#') or href.startswith('javascript:') or href.startswith('mailto:') or href.startswith('tel:'):
                return False
            # 协议相对路径 //cdn.xxx.com 不代理
            if href.startswith('//'):
                return False
            absolute = urljoin(base_url, href)
            parsed_href = urlparse(absolute)
            # 只代理同站点的HTML页面链接
            if parsed_href.netloc == parsed_base.netloc:
                return True
            # 外部链接也代理，确保在iframe内打开
            return True

        def make_proxy_url(target_url):
            """生成代理URL"""
            return f"/api/proxy/web?url={url_quote(target_url, safe='')}"

        # 1. 重写 <a> 标签 - 仅代理页面导航链接
        for a_tag in soup.find_all('a', href=True):
            original_href = a_tag['href']
            if not should_proxy(original_href):
                continue
            absolute_url = urljoin(base_url, original_href)
            a_tag['href'] = make_proxy_url(absolute_url)
            # 移除 target，确保在当前iframe中打开
            if 'target' in a_tag.attrs:
                del a_tag['target']

        # 2. 重写 <iframe> 标签
        for iframe_tag in soup.find_all('iframe', src=True):
            original_src = iframe_tag['src']
            if original_src.startswith('javascript:') or original_src.startswith('//'):
                continue
            absolute_url = urljoin(base_url, original_src)
            iframe_tag['src'] = make_proxy_url(absolute_url)

        # 3. 重写 <form> 标签 - 保留原始method，action通过代理
        for form_tag in soup.find_all('form'):
            action = form_tag.get('action', '')
            if action and not action.startswith('javascript:'):
                absolute_url = urljoin(base_url, action)
                form_tag['action'] = make_proxy_url(absolute_url)
            elif not action:
                form_tag['action'] = make_proxy_url(base_url)
            # 保留原始 method，不强制改为 GET
            # 添加隐藏字段传递目标URL（用于POST代理）
            hidden_input = soup.new_tag('input')
            hidden_input.attrs['type'] = 'hidden'
            hidden_input.attrs['name'] = 'url'
            hidden_input.attrs['value'] = absolute_url if action else base_url
            form_tag.append(hidden_input)
            # 将 form method 改为 POST 指向我们的代理
            form_tag['method'] = 'POST'
            form_tag['action'] = '/api/proxy/web'

        # 4. 重写 <meta http-equiv="refresh"> 跳转
        for meta_tag in soup.find_all('meta', attrs={'http-equiv': 'refresh'}):
            content = meta_tag.get('content', '')
            # 格式: "5;url=http://example.com"
            match = re.search(r'url=(.+)', content, re.IGNORECASE)
            if match:
                redirect_url = match.group(1).strip().strip("'\"")
                absolute_url = urljoin(base_url, redirect_url)
                meta_tag['content'] = content[:match.start(1)] + make_proxy_url(absolute_url)

        # 5. 重写 <area> 标签（图片地图链接）
        for area_tag in soup.find_all('area', href=True):
            original_href = area_tag['href']
            if not should_proxy(original_href):
                continue
            absolute_url = urljoin(base_url, original_href)
            area_tag['href'] = make_proxy_url(absolute_url)

        # 注意：不重写 <link>(CSS)、<script>(JS)、<img>(图片) 的 src/href
        # 这些静态资源直接使用原始URL，不经过代理，大幅提升加载速度

        # 6. 移除可能阻止iframe加载的 <meta> 标签
        for meta_tag in soup.find_all('meta'):
            http_equiv = meta_tag.get('http-equiv', '').lower()
            if http_equiv == 'x-frame-options' or http_equiv == 'content-security-policy':
                meta_tag.decompose()

        # 7. 注入CSS（最小化干预，保持原貌）
        inject_style = soup.new_tag('style')
        inject_style.string = """
            /* 代理注入样式 - 最小化干预 */
            img { max-width: 100% !important; height: auto !important; }
        """
        head = soup.find('head')
        if head:
            head.append(inject_style)

        html_content = str(soup)

        return Response(
            content=html_content,
            media_type='text/html; charset=utf-8',
            headers={
                "X-Frame-Options": "ALLOWALL",
                "Content-Security-Policy": "frame-ancestors *",
                "Access-Control-Allow-Origin": "*",
            }
        )

    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="网页加载超时，请稍后重试")
    except httpx.ConnectError:
        raise HTTPException(status_code=502, detail="无法连接到目标网页")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"代理网页时出错: {str(e)}")

@app.post("/api/proxy/web", summary="网页代理（POST）")
async def proxy_web_post(request: Request):
    """POST方式代理网页（用于表单提交如搜索）"""
    form_data = await request.form()
    target_url = form_data.get("url", "")
    if not target_url:
        raise HTTPException(status_code=400, detail="缺少url参数")
    
    # 收集表单中除 url 外的其他字段作为 POST 数据
    post_data = {k: v for k, v in form_data.items() if k != "url"}
    
    try:
        parsed = urlparse(target_url)
        if parsed.scheme not in ('http', 'https'):
            raise HTTPException(status_code=400, detail="仅支持 http/https 协议")
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=400, detail="无效的URL")
    
    try:
        async with httpx.AsyncClient(
            follow_redirects=True,
            timeout=20.0,
            headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
                'Accept-Encoding': 'identity',
                'Referer': target_url,
            }
        ) as client:
            if post_data:
                response = await client.post(target_url, data=post_data)
            else:
                response = await client.get(target_url)
        
        # 复用 GET 代理的 HTML 处理逻辑
        content_type = response.headers.get('content-type', 'text/html')
        if 'text/html' not in content_type:
            return Response(
                content=response.content,
                media_type=content_type,
                headers={
                    "X-Frame-Options": "ALLOWALL",
                    "Content-Security-Policy": "frame-ancestors *",
                    "Access-Control-Allow-Origin": "*",
                }
            )
        
        html_content = response.text
        soup = BeautifulSoup(html_content, 'html.parser')
        
        base_url = target_url
        parsed_base = urlparse(base_url)
        base_origin = f"{parsed_base.scheme}://{parsed_base.netloc}"
        
        def should_proxy(href):
            if not href:
                return False
            href = href.strip()
            if href.startswith('#') or href.startswith('javascript:') or href.startswith('mailto:') or href.startswith('tel:'):
                return False
            if href.startswith('//'):
                return False
            return True
        
        def make_proxy_url(t_url):
            return f"/api/proxy/web?url={url_quote(t_url, safe='')}"
        
        for a_tag in soup.find_all('a', href=True):
            original_href = a_tag['href']
            if not should_proxy(original_href):
                continue
            absolute_url = urljoin(base_url, original_href)
            a_tag['href'] = make_proxy_url(absolute_url)
            if 'target' in a_tag.attrs:
                del a_tag['target']
        
        for iframe_tag in soup.find_all('iframe', src=True):
            original_src = iframe_tag['src']
            if original_src.startswith('javascript:') or original_src.startswith('//'):
                continue
            absolute_url = urljoin(base_url, original_src)
            iframe_tag['src'] = make_proxy_url(absolute_url)
        
        for form_tag in soup.find_all('form'):
            action = form_tag.get('action', '')
            if action and not action.startswith('javascript:'):
                absolute_url = urljoin(base_url, action)
                form_tag['action'] = make_proxy_url(absolute_url)
            elif not action:
                form_tag['action'] = make_proxy_url(base_url)
        
        for meta_tag in soup.find_all('meta', attrs={'http-equiv': 'refresh'}):
            content = meta_tag.get('content', '')
            match = re.search(r'url=(.+)', content, re.IGNORECASE)
            if match:
                redirect_url = match.group(1).strip().strip("'\"")
                absolute_url = urljoin(base_url, redirect_url)
                meta_tag['content'] = content[:match.start(1)] + make_proxy_url(absolute_url)
        
        for area_tag in soup.find_all('area', href=True):
            original_href = area_tag['href']
            if not should_proxy(original_href):
                continue
            absolute_url = urljoin(base_url, original_href)
            area_tag['href'] = make_proxy_url(absolute_url)
        
        for meta_tag in soup.find_all('meta'):
            http_equiv = meta_tag.get('http-equiv', '').lower()
            if http_equiv == 'x-frame-options' or http_equiv == 'content-security-policy':
                meta_tag.decompose()
        
        inject_style = soup.new_tag('style')
        inject_style.string = """
            img { max-width: 100% !important; height: auto !important; }
        """
        head = soup.find('head')
        if head:
            head.append(inject_style)
        
        html_content = str(soup)
        
        return Response(
            content=html_content,
            media_type='text/html; charset=utf-8',
            headers={
                "X-Frame-Options": "ALLOWALL",
                "Content-Security-Policy": "frame-ancestors *",
                "Access-Control-Allow-Origin": "*",
            }
        )
    
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="网页加载超时，请稍后重试")
    except httpx.ConnectError:
        raise HTTPException(status_code=502, detail="无法连接到目标网页")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"代理网页失败: {str(e)}")


# ==================== 大模型配置路由 ====================

@app.get("/api/llm/configs", response_model=List[LLMConfigResponse], summary="获取所有大模型配置")
def get_llm_configs(db: Session = Depends(get_db)):
    return db.query(LLMConfig).order_by(LLMConfig.created_at.desc()).all()

@app.get("/api/llm/configs/active", summary="获取当前激活的大模型配置")
def get_active_llm_config(db: Session = Depends(get_db)):
    config = db.query(LLMConfig).filter(LLMConfig.is_active == True).first()
    if not config:
        raise HTTPException(status_code=404, detail="未配置大模型，请在设置中添加并激活")
    return config

@app.post("/api/llm/configs", response_model=LLMConfigResponse, summary="创建大模型配置")
def create_llm_config(config: LLMConfigCreate, db: Session = Depends(get_db)):
    # 如果设为激活，先取消其他激活状态
    if config.is_active:
        db.query(LLMConfig).update({LLMConfig.is_active: False})
    db_config = LLMConfig(**config.model_dump())
    db.add(db_config)
    db.commit()
    db.refresh(db_config)
    return db_config

@app.put("/api/llm/configs/{config_id}", response_model=LLMConfigResponse, summary="更新大模型配置")
def update_llm_config(config_id: int, config: LLMConfigUpdate, db: Session = Depends(get_db)):
    db_config = db.query(LLMConfig).filter(LLMConfig.id == config_id).first()
    if not db_config:
        raise HTTPException(status_code=404, detail="配置不存在")
    update_data = config.model_dump(exclude_unset=True)
    # 如果设为激活，先取消其他激活状态
    if update_data.get("is_active"):
        db.query(LLMConfig).filter(LLMConfig.id != config_id).update({LLMConfig.is_active: False})
    for key, value in update_data.items():
        setattr(db_config, key, value)
    db.commit()
    db.refresh(db_config)
    return db_config

@app.delete("/api/llm/configs/{config_id}", summary="删除大模型配置")
def delete_llm_config(config_id: int, db: Session = Depends(get_db)):
    db_config = db.query(LLMConfig).filter(LLMConfig.id == config_id).first()
    if not db_config:
        raise HTTPException(status_code=404, detail="配置不存在")
    db.delete(db_config)
    db.commit()
    return {"message": "配置已删除", "id": config_id}

@app.put("/api/llm/configs/{config_id}/activate", summary="激活大模型配置")
def activate_llm_config(config_id: int, db: Session = Depends(get_db)):
    db_config = db.query(LLMConfig).filter(LLMConfig.id == config_id).first()
    if not db_config:
        raise HTTPException(status_code=404, detail="配置不存在")
    db.query(LLMConfig).update({LLMConfig.is_active: False})
    db_config.is_active = True
    db.commit()
    db.refresh(db_config)
    return db_config


# ==================== 大模型调用核心函数 ====================

async def call_llm(messages: list, config: LLMConfig = None, max_tokens: int = None, temperature: float = None) -> str:
    """调用大模型，返回回复文本"""
    if not config:
        raise HTTPException(status_code=400, detail="未配置大模型")

    api_base = config.api_base.rstrip('/')
    url = f"{api_base}/chat/completions"

    headers = {"Content-Type": "application/json"}
    if config.api_key:
        headers["Authorization"] = f"Bearer {config.api_key}"

    payload = {
        "model": config.model,
        "messages": messages,
        "max_tokens": max_tokens or config.max_tokens,
        "temperature": temperature if temperature is not None else config.temperature,
    }

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.post(url, json=payload, headers=headers)
            response.raise_for_status()
            data = response.json()
            return data["choices"][0]["message"]["content"]
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="大模型响应超时，请稍后重试")
    except httpx.HTTPStatusError as e:
        raise HTTPException(status_code=502, detail=f"大模型API错误: {e.response.status_code}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"调用大模型失败: {str(e)}")


def get_active_config(db: Session) -> LLMConfig:
    """获取当前激活的大模型配置"""
    config = db.query(LLMConfig).filter(LLMConfig.is_active == True).first()
    if not config:
        raise HTTPException(status_code=400, detail="未配置大模型，请在设置中添加并激活一个配置")
    return config


# ==================== 大模型功能路由 ====================

@app.post("/api/llm/chat", summary="通用对话")
async def llm_chat(request: LLMChatRequest, db: Session = Depends(get_db)):
    """通用大模型对话"""
    config = get_active_config(db)
    content = await call_llm(request.messages, config, request.max_tokens, request.temperature)
    return LLMChatResponse(content=content, model=config.model, provider=config.provider)


@app.post("/api/llm/qa", summary="智能问答（知识库增强）")
async def llm_qa(request: LLMQARequest, db: Session = Depends(get_db)):
    """基于知识库的智能问答"""
    config = get_active_config(db)

    # 从知识库搜索相关内容
    keywords = [kw.strip() for kw in re.split(r'[\s,，。.!！?？;；:：、]+', request.question) if len(kw.strip()) >= 2]
    if not keywords:
        keywords = [request.question.strip()]

    knowledge_items = []
    for kw in keywords[:5]:
        items = db.query(KnowledgeItem).filter(
            or_(
                KnowledgeItem.title.contains(kw),
                KnowledgeItem.content.contains(kw),
            )
        ).limit(3).all()
        knowledge_items.extend(items)

    # 去重
    seen = set()
    unique_items = []
    for item in knowledge_items:
        if item.id not in seen:
            seen.add(item.id)
            unique_items.append(item)

    # 构建提示词
    context = "\n".join([f"- {item.title}: {item.content[:500]}" for item in unique_items[:10]])

    messages = [
        {"role": "system", "content": "你是一个专业的知识助手。请根据以下知识库内容回答用户的问题。如果知识库中没有相关信息，请根据你的知识回答，并说明这不是来自知识库。\n\n知识库内容：\n" + context},
        {"role": "user", "content": request.question}
    ]

    content = await call_llm(messages, config)
    return {
        "answer": content,
        "sources": [{"id": item.id, "title": item.title} for item in unique_items[:5]],
        "model": config.model,
    }


@app.post("/api/llm/writing", summary="写作助手增强")
async def llm_writing(request: LLMWritingRequest, db: Session = Depends(get_db)):
    """写作助手：润色/续写/总结/扩写"""
    config = get_active_config(db)

    prompts = {
        "polish": f"请润色以下文稿，改善语言表达，修正语病，保持原意不变，直接返回润色后的内容：\n\n{request.content}",
        "continue": f"请根据以下文稿的上下文和风格，续写约200字，直接返回续写内容：\n\n{request.content}",
        "summarize": f"请总结以下文稿的主要内容，提炼要点，直接返回总结：\n\n{request.content}",
        "expand": f"请扩写以下内容，丰富细节和论述，保持原有风格，直接返回扩写后的内容：\n\n{request.content}",
    }

    instruction = request.instruction.strip() if request.instruction else ""
    user_content = prompts.get(request.action, f"{instruction}\n\n{request.content}" if instruction else request.content)

    messages = [
        {"role": "system", "content": "你是一个专业的中文写作助手。请直接返回处理后的内容，不要添加额外说明。"},
        {"role": "user", "content": user_content}
    ]

    content = await call_llm(messages, config, max_tokens=4096)
    return {"content": content, "action": request.action}


@app.post("/api/llm/compose", summary="AI写作（素材+提纲）")
async def llm_compose(request_body: dict, db: Session = Depends(get_db)):
    """基于知识库素材和用户提纲进行AI写作
    - **material_ids**: 知识库素材ID列表
    - **outline**: 用户提供的写作提纲/要求
    - **style**: 写作风格（formal/casual/academic/news）
    - **length**: 期望字数（short/medium/long）
    """
    config = get_active_config(db)
    
    material_ids = request_body.get("material_ids", [])
    outline = request_body.get("outline", "").strip()
    style = request_body.get("style", "formal")
    length = request_body.get("length", "medium")
    
    if not outline:
        raise HTTPException(status_code=400, detail="请提供写作提纲")
    
    # 获取素材内容
    materials = []
    if material_ids:
        items = db.query(KnowledgeItem).filter(KnowledgeItem.id.in_(material_ids)).all()
        for item in items:
            materials.append(f"【{item.title}】{item.content[:800]}")
    
    # 风格和长度映射
    style_map = {
        "formal": "正式公文风格，语言规范严谨",
        "casual": "通俗易懂风格，语言自然流畅",
        "academic": "学术论文风格，论述严谨有据",
        "news": "新闻报道风格，客观简洁明了",
    }
    length_map = {
        "short": "500字左右",
        "medium": "1000-1500字",
        "long": "2000字以上",
    }
    
    style_desc = style_map.get(style, "正式公文风格")
    length_desc = length_map.get(length, "1000-1500字")
    
    # 构建提示词
    material_text = "\n\n".join(materials) if materials else "（无参考素材）"
    
    messages = [
        {"role": "system", "content": f"""你是一位专业的中文写作助手。请根据用户提供的参考素材和写作提纲，撰写一篇高质量的文稿。

写作要求：
- 风格：{style_desc}
- 篇幅：{length_desc}
- 充分利用参考素材中的信息和观点
- 内容要有逻辑性、条理清晰
- 语言流畅，用词准确
- 直接输出文稿内容，不要添加额外说明或标题标记"""},
        {"role": "user", "content": f"""参考素材：
{material_text}

写作提纲/要求：
{outline}

请根据以上素材和提纲撰写文稿。"""}
    ]
    
    content = await call_llm(messages, config, max_tokens=4096, temperature=0.7)
    return {"content": content, "material_count": len(materials), "style": style, "length": length}


@app.post("/api/llm/knowledge", summary="知识库智能处理")
async def llm_knowledge(request: LLMKnowledgeRequest, db: Session = Depends(get_db)):
    """知识库智能处理：提取/分类/总结"""
    config = get_active_config(db)

    prompts = {
        "extract": "请从以下内容中提取关键知识点，每行一个，格式为「知识点：解释」：\n\n",
        "classify": "请为以下内容推荐合适的分类和标签，返回JSON格式 {\"category\": \"分类名\", \"tags\": [\"标签1\", \"标签2\"]}：\n\n",
        "summarize": "请用简洁的语言总结以下内容的要点：\n\n",
    }

    user_content = prompts.get(request.action, "") + request.content
    messages = [
        {"role": "system", "content": "你是一个专业的知识管理助手。请直接返回结果，不要添加额外说明。"},
        {"role": "user", "content": user_content}
    ]

    content = await call_llm(messages, config)
    return {"content": content, "action": request.action}


@app.post("/api/llm/proofread", summary="智能核稿")
async def llm_proofread(request: LLMProofreadRequest, db: Session = Depends(get_db)):
    """大模型辅助核稿"""
    config = get_active_config(db)

    messages = [
        {"role": "system", "content": """你是一个专业的中文文稿核稿助手。请检查以下文稿中的问题，包括：
1. 错别字和用词错误
2. 标点符号错误
3. 语法问题
4. 逻辑不通顺的地方
5. 格式问题

请以JSON数组格式返回发现的问题，每个问题包含：
- "type": 问题类型（错别字/标点/语法/逻辑/格式）
- "text": 有问题的原文片段
- "suggestion": 修改建议
- "severity": 严重程度（error/warning/info）

只返回JSON数组，不要其他内容。如果没有发现问题，返回空数组 []。"""},
        {"role": "user", "content": request.content}
    ]

    content = await call_llm(messages, config, max_tokens=4096, temperature=0.3)

    # 尝试解析JSON
    try:
        # 提取JSON部分
        json_match = re.search(r'\[.*\]', content, re.DOTALL)
        if json_match:
            issues = json.loads(json_match.group())
        else:
            issues = []
    except json.JSONDecodeError:
        issues = []

    return {"issues": issues, "raw_response": content}


# ==================== 静态文件服务（生产模式） ====================

static_dir = os.path.join(os.path.dirname(__file__), "static")

@app.get("/", include_in_schema=False)
async def serve_root():
    """根路径：优先返回前端页面，否则返回API信息"""
    if os.path.isdir(static_dir):
        index_path = os.path.join(static_dir, "index.html")
        if os.path.isfile(index_path):
            return FileResponse(index_path)
    return {
        "name": "知识管理系统 API",
        "version": "1.0.0",
        "docs": "/docs",
        "redoc": "/redoc",
    }

if os.path.isdir(static_dir):
    @app.get("/{full_path:path}", include_in_schema=False)
    async def serve_static(full_path: str):
        """生产模式下提供前端静态文件"""
        file_path = os.path.join(static_dir, full_path)
        if os.path.isfile(file_path):
            return FileResponse(file_path)
        # SPA 路由回退：返回 index.html
        index_path = os.path.join(static_dir, "index.html")
        if os.path.isfile(index_path):
            return FileResponse(index_path)
        raise HTTPException(status_code=404, detail="文件不存在")
